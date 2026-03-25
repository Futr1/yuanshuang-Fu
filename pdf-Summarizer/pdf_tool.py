import sys
import os
import re


def extract(pdf_path, out_path):
    """将 PDF 文本直接写入文件，完全绕过终端编码限制。"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        text = "\n".join(pages_text)
        # 清理 PDF 解码时产生的孤立代理字符（\ud800-\udfff），
        # 这类字符来自 PDF 内部编码缺陷，UTF-8 无法写入
        text = re.sub(r'[\ud800-\udfff]', '', text)
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"提取成功，共 {len(reader.pages)} 页 → {out_path}")
    except Exception as e:
        print(f"读取 PDF 失败: {e}")
        sys.exit(1)


def _add_run_with_inline_bold(paragraph, line):
    """解析行内 **粗体** 标记并分段写入 run。"""
    parts = re.split(r'\*\*(.+?)\*\*', line)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def _is_table_row(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and len(s) > 2


def _is_separator_row(line):
    """判断是否为 |---|---| 这样的分隔行。"""
    s = line.strip()
    return _is_table_row(s) and bool(re.match(r'^[|\-:= ]+$', s))


def _parse_cells(line):
    """把 | a | b | c | 解析为 ['a', 'b', 'c']。"""
    parts = line.strip().strip('|').split('|')
    return [p.strip() for p in parts]


def _render_table_block(doc, table_lines):
    """将收集到的 Markdown 表格行渲染为 Word 表格。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    data_rows = [_parse_cells(r) for r in table_lines if not _is_separator_row(r)]
    if not data_rows:
        return

    ncols = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=ncols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data_rows):
        for j in range(ncols):
            cell_text = row_data[j] if j < len(row_data) else ''
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            # 设置单元格字体
            run = p.add_run(cell_text)
            run.font.name = '微软雅黑'
            run.font.size = Pt(10.5)
            # 表头行加粗
            if i == 0:
                run.bold = True
            # 东亚字体
            rPr = run._r.get_or_add_rPr()
            from docx.oxml import OxmlElement
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '微软雅黑')


def save(txt_path, docx_path):
    """将摘要文本转换为格式良好的 Word 文档。
    支持：# 标题层级、**粗体**、- 列表、Markdown 表格。
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()

        doc = Document()

        # ── 全局默认字体：微软雅黑 ────────────────────────────────
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # ── 逐行解析，用状态机处理表格块 ─────────────────────────
        lines = content.split('\n')
        table_buffer = []   # 收集当前表格的行

        def flush_table():
            if table_buffer:
                _render_table_block(doc, table_buffer)
                table_buffer.clear()

        for line in lines:
            stripped = line.strip()

            # ── 表格行：先缓冲，遇到非表格行时整块渲染 ──────────
            if _is_table_row(stripped):
                table_buffer.append(stripped)
                continue

            # 遇到非表格行，先把缓冲的表格刷出来
            flush_table()

            if not stripped:
                continue

            # ── 标题层级 ──────────────────────────────────────────
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            # ── 无序列表 ──────────────────────────────────────────
            elif re.match(r'^[-*•]\s', stripped):
                p = doc.add_paragraph(style='List Bullet')
                _add_run_with_inline_bold(p, stripped[2:])
            # ── 有序列表 ──────────────────────────────────────────
            elif re.match(r'^\d+\.\s', stripped):
                p = doc.add_paragraph(style='List Number')
                _add_run_with_inline_bold(p, re.sub(r'^\d+\.\s', '', stripped))
            # ── 整行粗体 ──────────────────────────────────────────
            elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:-2])
                run.bold = True
            # ── 普通段落（可含行内粗体） ──────────────────────────
            else:
                p = doc.add_paragraph()
                _add_run_with_inline_bold(p, stripped)

        # 文件末尾可能还有未刷出的表格
        flush_table()

        doc.save(docx_path)
        print(f"成功将摘要保存至: {docx_path}")
    except Exception as e:
        print(f"保存 DOCX 失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    action = sys.argv[1] if len(sys.argv) > 1 else ''

    if action == "extract" and len(sys.argv) == 4:
        extract(sys.argv[2], sys.argv[3])
    elif action == "save" and len(sys.argv) == 4:
        save(sys.argv[2], sys.argv[3])
    else:
        print("用法:")
        print("  提取: python pdf_tool.py extract <pdf路径> <输出txt路径>")
        print("  保存: python pdf_tool.py save   <txt路径> <输出docx路径>")
        sys.exit(1)
